"""Word document generation tool.

Generates a .docx file from markdown-ish text and returns a download link.
The file is saved to a configurable output directory (CORECODER_OUTPUT_DIR,
default ~/.corecoder/docs/) and served by the web server's /download/ route.
"""

import os
import re
from pathlib import Path

from .base import Tool

OUTPUT_DIR = Path(os.getenv("CORECODER_OUTPUT_DIR", str(Path.home() / ".corecoder" / "docs")))

# sanitize filename: keep alnum, dash, underscore only
_SAFE_NAME = re.compile(r"[^A-Za-z0-9_-]+")


class WriteDocxTool(Tool):
    name = "write_docx"
    description = (
        "Write content to a Word (.docx) document and return a download link. "
        "Use this when the user asks for a document, report, or summary file. "
        "Content supports markdown-style headings (#, ##, ###), "
        "bullet points (- or *), and numbered lists (1. 2. 3.)."
    )
    parameters = {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "File name without .docx extension (e.g. 'project-report')",
            },
            "title": {
                "type": "string",
                "description": "Document title, shown as the top-level heading",
            },
            "content": {
                "type": "string",
                "description": (
                    "Document body in markdown-ish format. "
                    "Supports: # H1, ## H2, ### H3, - bullet, 1. numbered, plain text."
                ),
            },
        },
        "required": ["filename", "title", "content"],
    }

    def execute(self, filename: str, title: str, content: str) -> str:
        try:
            from docx import Document
        except ImportError:
            return "Error: python-docx not installed. Run: pip install python-docx"

        safe_name = _SAFE_NAME.sub("-", filename).strip("-_") or "document"
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        path = OUTPUT_DIR / f"{safe_name}.docx"
        # avoid overwriting: append -1, -2, ... if file exists
        counter = 1
        while path.exists():
            path = OUTPUT_DIR / f"{safe_name}-{counter}.docx"
            counter += 1

        doc = Document()
        doc.add_heading(title, level=0)

        for line in content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", stripped), style="List Number")
            else:
                doc.add_paragraph(stripped)

        doc.save(path)
        return f"文档已生成。下载链接：/download/{path.name}"
